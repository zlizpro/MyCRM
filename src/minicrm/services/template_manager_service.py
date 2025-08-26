"""
MiniCRM 模板管理服务

专门负责文档模板的管理,包括:
- 模板创建和更新
- 模板删除和查询
- 模板内容管理
- 模板验证

设计原则:
- 单一职责:只负责模板管理
- 支持多种模板格式
- 完整的错误处理和日志记录
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any


class TemplateManagerService:
    """
    模板管理服务

    负责管理各种文档模板的创建、更新、删除和查询.
    """

    def __init__(self):
        """初始化模板管理服务"""
        self._logger = logging.getLogger(__name__)

        # 模板文件路径
        self._template_dir = Path(__file__).parent.parent / "resources" / "templates"

        # 支持的模板类型
        self._template_types = {
            "contract": "contract_template.docx",
            "quote": "quote_template.docx",
            "customer_report": "customer_report_template.docx",
        }

        # 检查模板目录
        if not self._template_dir.exists():
            self._template_dir.mkdir(parents=True, exist_ok=True)

        self._logger.info("模板管理服务初始化完成")

    def get_available_templates(self) -> list[str]:
        """
        获取可用的模板列表

        Returns:
            List[str]: 可用模板类型列表
        """
        return list(self._template_types.keys())

    def get_template_path(self, template_type: str) -> Path | None:
        """
        获取模板文件路径

        Args:
            template_type: 模板类型

        Returns:
            Path | None: 模板文件路径,不存在则返回None
        """
        if template_type not in self._template_types:
            return None

        template_path = self._template_dir / self._template_types[template_type]
        return template_path if template_path.exists() else None

    def create_custom_template(
        self, template_name: str, template_content: str, template_type: str = "custom"
    ) -> bool:
        """
        创建增强的自定义模板 - 优化版本

        支持多种模板格式和高级功能:
        - 变量占位符支持
        - 条件内容块
        - 循环数据块
        - 样式和格式设置
        - 模板验证和预览

        Args:
            template_name: 模板名称
            template_content: 模板内容(支持Jinja2语法)
            template_type: 模板类型

        Returns:
            bool: 创建是否成功
        """
        try:
            self._logger.info(f"创建增强自定义模板: {template_name}")

            # 创建自定义模板目录结构
            custom_template_dir = self._template_dir / "custom"
            custom_template_dir.mkdir(exist_ok=True)

            # 创建模板元数据目录
            metadata_dir = custom_template_dir / "metadata"
            metadata_dir.mkdir(exist_ok=True)

            # 验证模板内容
            validation_result = self._validate_template_content(template_content)
            if not validation_result["valid"]:
                self._logger.warning(f"模板内容验证失败: {validation_result['errors']}")

            # 保存模板文件
            template_file = custom_template_dir / f"{template_name}.docx"

            # 创建增强的Word文档模板
            success = self._create_enhanced_word_template(
                template_file, template_name, template_content, template_type
            )

            if success:
                # 保存模板元数据
                self._save_template_metadata(
                    template_name,
                    {
                        "type": template_type,
                        "content": template_content,
                        "variables": validation_result.get("variables", []),
                        "created_at": self._get_current_timestamp(),
                        "version": "1.0",
                    },
                )

                # 更新模板类型映射
                self._template_types[template_name] = f"custom/{template_name}.docx"
                self._logger.info(f"增强自定义模板创建成功: {template_file}")
                return True
            else:
                # 使用增强文本模板作为备用
                return self._create_enhanced_text_template(
                    template_name, template_content, template_type
                )

        except Exception as e:
            self._logger.error(f"创建增强自定义模板失败: {e}")
            return False

    def _create_word_template(
        self, template_file: Path, template_name: str, content: str
    ) -> bool:
        """创建Word模板"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(f"{template_name}模板", 0)

            # 添加模板内容
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

            doc.save(str(template_file))
            return True

        except ImportError:
            self._logger.warning("python-docx库未安装,将使用文本模板")
            return False
        except Exception as e:
            self._logger.error(f"创建Word模板失败: {e}")
            return False

    def _create_text_template(self, template_name: str, template_content: str) -> bool:
        """创建文本模板(备用方案)"""
        try:
            custom_template_dir = self._template_dir / "custom"
            custom_template_dir.mkdir(exist_ok=True)

            template_file = custom_template_dir / f"{template_name}.txt"

            with open(template_file, "w", encoding="utf-8") as f:
                f.write(f"{template_name}模板\n")
                f.write("=" * 20 + "\n\n")
                f.write(template_content)

            self._template_types[template_name] = f"custom/{template_name}.txt"
            return True

        except Exception as e:
            self._logger.error(f"创建文本模板失败: {e}")
            return False

    def update_template(self, template_name: str, template_content: str) -> bool:
        """
        更新现有模板

        Args:
            template_name: 模板名称
            template_content: 新的模板内容

        Returns:
            bool: 更新是否成功
        """
        try:
            if template_name not in self._template_types:
                return self.create_custom_template(template_name, template_content)

            template_file_path = (
                self._template_dir / self._template_types[template_name]
            )

            if template_file_path.suffix == ".docx":
                return self._update_word_template(template_file_path, template_content)
            else:
                return self._update_text_template(template_file_path, template_content)

        except Exception as e:
            self._logger.error(f"更新模板失败: {e}")
            return False

    def _update_word_template(self, template_file: Path, content: str) -> bool:
        """更新Word模板"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("更新的模板", 0)

            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

            doc.save(str(template_file))
            return True

        except ImportError:
            return self._update_text_template(
                template_file.with_suffix(".txt"), content
            )
        except Exception as e:
            self._logger.error(f"更新Word模板失败: {e}")
            return False

    def _update_text_template(self, template_file: Path, content: str) -> bool:
        """更新文本模板"""
        try:
            with open(template_file, "w", encoding="utf-8") as f:
                f.write("更新的模板\n")
                f.write("=" * 20 + "\n\n")
                f.write(content)
            return True

        except Exception as e:
            self._logger.error(f"更新文本模板失败: {e}")
            return False

    def delete_template(self, template_name: str) -> bool:
        """
        删除模板

        Args:
            template_name: 模板名称

        Returns:
            bool: 删除是否成功
        """
        try:
            if template_name not in self._template_types:
                return False

            template_file_path = (
                self._template_dir / self._template_types[template_name]
            )

            if template_file_path.exists():
                template_file_path.unlink()

            # 从模板类型映射中移除
            del self._template_types[template_name]

            self._logger.info(f"模板删除成功: {template_name}")
            return True

        except Exception as e:
            self._logger.error(f"删除模板失败: {e}")
            return False

    def get_template_content(self, template_name: str) -> str:
        """
        获取模板内容

        Args:
            template_name: 模板名称

        Returns:
            str: 模板内容
        """
        try:
            if template_name not in self._template_types:
                return ""

            template_file_path = (
                self._template_dir / self._template_types[template_name]
            )

            if not template_file_path.exists():
                return ""

            if template_file_path.suffix == ".txt":
                with open(template_file_path, encoding="utf-8") as f:
                    return f.read()
            else:
                # 对于Word文档,返回简化的内容描述
                return f"Word模板文件: {template_file_path.name}"

        except Exception as e:
            self._logger.error(f"获取模板内容失败: {e}")
            return ""

    def is_valid_word_template(self, template_file: Path) -> bool:
        """
        检查是否为有效的Word模板文件

        Args:
            template_file: 模板文件路径

        Returns:
            bool: 是否为有效的Word模板
        """
        try:
            # 检查文件扩展名
            if template_file.suffix.lower() != ".docx":
                return False

            # 检查文件大小(空文件或过小的文件可能无效)
            if template_file.stat().st_size < 1000:  # 小于1KB可能是无效文件
                return False

            # 尝试读取文件头部验证格式
            with open(template_file, "rb") as f:
                header = f.read(4)
                # Word文档应该以PK开头(ZIP格式)
                if not header.startswith(b"PK"):
                    return False

            return True

        except Exception:
            return False

    def _validate_template_content(self, content: str) -> dict[str, Any]:
        """
        验证模板内容

        检查模板语法、变量占位符等

        Args:
            content: 模板内容

        Returns:
            Dict[str, Any]: 验证结果
        """

        result: dict[str, Any] = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "variables": [],
            "blocks": [],
        }

        try:
            # 提取变量占位符 {{variable_name}}
            variable_pattern = r"\{\{\s*(\w+)\s*\}\}"
            variables = re.findall(variable_pattern, content)
            result["variables"] = list(set(variables))

            # 检查条件块 {% if condition %}...{% endif %}
            if_pattern = r"\{\%\s*if\s+.*?\%\}.*?\{\%\s*endif\s*\%\}"
            if_blocks = re.findall(if_pattern, content, re.DOTALL)
            result["blocks"].extend(if_blocks)

            # 检查循环块 {% for item in items %}...{% endfor %}
            for_pattern = r"\{\%\s*for\s+.*?\%\}.*?\{\%\s*endfor\s*\%\}"
            for_blocks = re.findall(for_pattern, content, re.DOTALL)
            result["blocks"].extend(for_blocks)

            # 检查语法错误
            if content.count("{{") != content.count("}}"):
                result["errors"].append("变量占位符不匹配")
                result["valid"] = False

            if content.count("{%") != content.count("%}"):
                result["errors"].append("控制块标签不匹配")
                result["valid"] = False

            # 检查常见问题
            if len(variables) == 0:
                result["warnings"].append("模板中未发现变量占位符")

        except Exception as e:
            result["valid"] = False
            result["errors"].append(f"模板验证异常: {e}")

        return result

    def _create_enhanced_word_template(
        self, template_file: Path, template_name: str, content: str, template_type: str
    ) -> bool:
        """创建增强的Word模板"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor

            doc = Document()

            # 设置文档属性
            doc.core_properties.title = f"{template_name}模板"
            doc.core_properties.author = "MiniCRM系统"
            doc.core_properties.subject = f"{template_type}类型模板"

            # 添加标题
            title = doc.add_heading(f"{template_name}模板", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加模板说明
            doc.add_heading("模板说明", level=1)
            doc.add_paragraph(
                f"这是一个{template_type}类型的文档模板,支持变量替换和动态内容生成."
            )

            # 添加变量说明
            validation_result = self._validate_template_content(content)
            variables = validation_result.get("variables", [])

            if variables:
                doc.add_heading("可用变量", level=2)
                var_table = doc.add_table(rows=1, cols=2)
                var_table.style = "Table Grid"

                # 表头
                hdr_cells = var_table.rows[0].cells
                hdr_cells[0].text = "变量名"
                hdr_cells[1].text = "说明"

                # 变量列表
                for var in variables:
                    row_cells = var_table.add_row().cells
                    row_cells[0].text = f"{{{{{var}}}}}"
                    row_cells[1].text = f"{var}的值"

            # 添加模板内容
            doc.add_heading("模板内容", level=1)

            # 处理模板内容,按段落分割
            for line in content.split("\n"):
                if line.strip():
                    paragraph = doc.add_paragraph(line)
                    # 为包含变量的段落添加特殊样式
                    if "{{" in line and "}}" in line:
                        for run in paragraph.runs:
                            if "{{" in run.text:
                                run.font.color.rgb = RGBColor(0, 100, 200)  # 蓝色
                                run.bold = True

            # 添加页脚
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"由MiniCRM系统生成 - {template_name}模板"

            doc.save(str(template_file))
            return True

        except ImportError:
            self._logger.warning("python-docx库未安装,将使用增强文本模板")
            return False
        except Exception as e:
            self._logger.error(f"创建增强Word模板失败: {e}")
            return False

    def _create_enhanced_text_template(
        self, template_name: str, template_content: str, template_type: str
    ) -> bool:
        """创建增强的文本模板(备用方案)"""
        try:
            custom_template_dir = self._template_dir / "custom"
            custom_template_dir.mkdir(exist_ok=True)

            template_file = custom_template_dir / f"{template_name}.txt"

            # 生成增强的文本模板
            enhanced_content = self._generate_enhanced_text_template(
                template_name, template_content, template_type
            )

            with open(template_file, "w", encoding="utf-8") as f:
                f.write(enhanced_content)

            self._template_types[template_name] = f"custom/{template_name}.txt"
            return True

        except Exception as e:
            self._logger.error(f"创建增强文本模板失败: {e}")
            return False

    def _generate_enhanced_text_template(
        self, template_name: str, content: str, template_type: str
    ) -> str:
        """生成增强的文本模板内容"""
        validation_result = self._validate_template_content(content)
        variables = validation_result.get("variables", [])

        enhanced_content = f"""
{template_name}模板
{"=" * (len(template_name) + 2)}

模板类型: {template_type}
创建时间: {self._get_current_timestamp()}
系统版本: MiniCRM v1.0

模板说明:
---------
这是一个{template_type}类型的文档模板,支持变量替换和动态内容生成.

可用变量:
---------
"""

        if variables:
            for var in variables:
                enhanced_content += f"- {{{{{var}}}}} : {var}的值\n"
        else:
            enhanced_content += "- 无变量占位符\n"

        enhanced_content += f"""

模板内容:
---------
{content}

使用说明:
---------
1. 变量使用双大括号包围,如: {{{{customer_name}}}}
2. 支持条件语句: {{% if condition %}}...{{% endif %}}
3. 支持循环语句: {{% for item in items %}}...{{% endfor %}}
4. 生成文档时,变量将被实际数据替换

注意事项:
---------
- 请确保变量名称正确
- 条件和循环语句需要正确闭合
- 建议在使用前进行模板预览测试
"""

        return enhanced_content

    def _save_template_metadata(
        self, template_name: str, metadata: dict[str, Any]
    ) -> bool:
        """保存模板元数据"""
        try:
            import json

            metadata_dir = self._template_dir / "custom" / "metadata"
            metadata_file = metadata_dir / f"{template_name}.json"

            with open(metadata_file, "w", encoding="utf-8") as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)

            return True

        except Exception as e:
            self._logger.error(f"保存模板元数据失败: {e}")
            return False

    def _get_current_timestamp(self) -> str:
        """获取当前时间戳"""

        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def get_template_variables(self, template_name: str) -> list[str]:
        """
        获取模板中的变量列表

        Args:
            template_name: 模板名称

        Returns:
            List[str]: 变量列表
        """
        try:
            content = self.get_template_content(template_name)
            if content:
                validation_result = self._validate_template_content(content)
                return validation_result.get("variables", [])
            return []

        except Exception as e:
            self._logger.error(f"获取模板变量失败: {e}")
            return []

    def preview_template(self, template_name: str, sample_data: dict[str, Any]) -> str:
        """
        预览模板效果

        Args:
            template_name: 模板名称
            sample_data: 示例数据

        Returns:
            str: 预览内容
        """
        try:
            content = self.get_template_content(template_name)
            if not content:
                return "模板内容为空"

            # 简单的变量替换预览
            preview_content = content
            for key, value in sample_data.items():
                placeholder = f"{{{{{key}}}}}"
                preview_content = preview_content.replace(placeholder, str(value))

            return preview_content

        except Exception as e:
            self._logger.error(f"模板预览失败: {e}")
            return f"预览失败: {e}"

    def duplicate_template(self, source_template: str, new_template_name: str) -> bool:
        """
        复制模板

        Args:
            source_template: 源模板名称
            new_template_name: 新模板名称

        Returns:
            bool: 复制是否成功
        """
        try:
            if source_template not in self._template_types:
                return False

            # 获取源模板内容
            source_content = self.get_template_content(source_template)
            if not source_content:
                return False

            # 创建新模板
            return self.create_custom_template(
                new_template_name, source_content, "custom"
            )

        except Exception as e:
            self._logger.error(f"复制模板失败: {e}")
            return False

    def get_template_info(self, template_name: str) -> dict[str, Any]:
        """
        获取增强的模板信息

        Args:
            template_name: 模板名称

        Returns:
            Dict[str, Any]: 模板信息字典
        """
        try:
            if template_name not in self._template_types:
                return {}

            template_file_path = (
                self._template_dir / self._template_types[template_name]
            )

            info = {
                "name": template_name,
                "file_path": str(template_file_path),
                "exists": template_file_path.exists(),
                "type": template_file_path.suffix,
                "variables": self.get_template_variables(template_name),
            }

            if template_file_path.exists():
                stat = template_file_path.stat()
                info.update(
                    {
                        "size": stat.st_size,
                        "modified": stat.st_mtime,
                        "is_valid": self.is_valid_word_template(template_file_path)
                        if template_file_path.suffix == ".docx"
                        else True,
                    }
                )

            # 尝试加载元数据
            try:
                import json

                metadata_file = (
                    self._template_dir / "custom" / "metadata" / f"{template_name}.json"
                )
                if metadata_file.exists():
                    with open(metadata_file, encoding="utf-8") as f:
                        metadata = json.load(f)
                        info.update(metadata)
            except Exception:
                pass  # 元数据不存在或损坏,忽略

            return info

        except Exception as e:
            self._logger.error(f"获取模板信息失败: {e}")
            return {}
