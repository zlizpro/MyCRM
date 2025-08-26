"""
MiniCRM 数据导入导出服务

提供统一的数据导入导出功能,包括:
- CSV/Excel数据导入导出
- Word/PDF文档生成
- 数据验证和格式转换
- 批量处理和进度跟踪

设计原则:
- 遵循MiniCRM分层架构标准
- 使用transfunctions进行数据验证和格式化
- 支持多种文件格式和数据类型
- 提供完整的错误处理和日志记录
"""

import csv
import logging
import os
from pathlib import Path
from typing import Any

from minicrm.core.exceptions import ServiceError
from minicrm.services.contract_service import ContractService
from minicrm.services.customer_service import CustomerService
from minicrm.services.supplier_service import SupplierService
from transfunctions import (
    validate_customer_data,
    validate_supplier_data,
)


class ImportExportService:
    """
    数据导入导出服务

    提供统一的数据导入导出功能,支持多种格式和数据类型.
    """

    def __init__(
        self,
        customer_service: CustomerService,
        supplier_service: SupplierService,
        contract_service: ContractService,
    ):
        """
        初始化导入导出服务

        Args:
            customer_service: 客户服务实例
            supplier_service: 供应商服务实例
            contract_service: 合同服务实例
        """
        self._customer_service = customer_service
        self._supplier_service = supplier_service
        self._contract_service = contract_service
        self._logger = logging.getLogger(__name__)

        # 延迟初始化可选服务
        self._pdf_service = None
        self._document_service = None

        # 支持的文件格式
        self._supported_import_formats = [".csv", ".xlsx"]
        self._supported_export_formats = [".csv", ".xlsx", ".pdf", ".docx"]

        # 数据类型映射
        self._data_type_services = {
            "customers": self._customer_service,
            "suppliers": self._supplier_service,
            "contracts": self._contract_service,
        }

        self._logger.info("导入导出服务初始化完成")

    def get_supported_formats(self) -> dict[str, list[str]]:
        """
        获取支持的文件格式

        Returns:
            Dict[str, List[str]]: 支持的导入和导出格式
        """
        return {
            "import": self._supported_import_formats,
            "export": self._supported_export_formats,
        }

    def validate_import_file(self, file_path: str) -> tuple[bool, str]:
        """
        验证导入文件

        Args:
            file_path: 文件路径

        Returns:
            Tuple[bool, str]: (是否有效, 错误信息)
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return False, "文件不存在"

            # 检查文件扩展名
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self._supported_import_formats:
                return False, f"不支持的文件格式: {file_ext}"

            # 检查文件大小(限制为50MB)
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                return False, f"文件过大,最大支持 {max_size // (1024 * 1024)}MB"

            # 尝试读取文件头部验证格式
            if file_ext == ".csv":
                return self._validate_csv_file(file_path)
            elif file_ext == ".xlsx":
                return self._validate_excel_file(file_path)

            return True, ""

        except Exception as e:
            self._logger.error(f"验证导入文件失败: {e}")
            return False, f"文件验证失败: {e}"

    def _validate_csv_file(self, file_path: str) -> tuple[bool, str]:
        """验证CSV文件"""
        try:
            with open(file_path, encoding="utf-8-sig") as file:
                # 读取前几行检查格式
                sample = file.read(1024)
                if not sample.strip():
                    return False, "文件为空"

                # 检查是否包含有效的CSV数据
                file.seek(0)
                reader = csv.reader(file)
                headers = next(reader, None)
                if not headers:
                    return False, "CSV文件没有标题行"

            return True, ""

        except UnicodeDecodeError:
            return False, "文件编码错误,请使用UTF-8编码"
        except Exception as e:
            return False, f"CSV文件格式错误: {e}"

    def _validate_excel_file(self, file_path: str) -> tuple[bool, str]:
        """验证Excel文件"""
        try:
            import pandas as pd

            # 尝试读取Excel文件
            df = pd.read_excel(file_path, nrows=0)  # 只读取标题行
            if df.empty:
                return False, "Excel文件为空"

            return True, ""

        except ImportError:
            return False, "需要安装pandas和openpyxl库来处理Excel文件"
        except Exception as e:
            return False, f"Excel文件格式错误: {e}"

    def preview_import_data(
        self, file_path: str, max_rows: int = 5
    ) -> tuple[list[str], list[dict[str, Any]]]:
        """
        预览导入数据

        Args:
            file_path: 文件路径
            max_rows: 最大预览行数

        Returns:
            Tuple[List[str], List[Dict[str, Any]]]: (字段列表, 数据行列表)

        Raises:
            ServiceError: 预览失败时抛出
        """
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext == ".csv":
                return self._preview_csv_data(file_path, max_rows)
            elif file_ext == ".xlsx":
                return self._preview_excel_data(file_path, max_rows)
            else:
                raise ServiceError(f"不支持预览的文件格式: {file_ext}")

        except Exception as e:
            self._logger.error(f"预览导入数据失败: {e}")
            raise ServiceError(f"预览导入数据失败: {e}") from e

    def _preview_csv_data(
        self, file_path: str, max_rows: int
    ) -> tuple[list[str], list[dict[str, Any]]]:
        """预览CSV数据"""
        with open(file_path, encoding="utf-8-sig") as file:
            # 自动检测分隔符
            sample = file.read(1024)
            file.seek(0)
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter

            reader = csv.DictReader(file, delimiter=delimiter)
            headers = list(reader.fieldnames or [])

            data_rows = []
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                data_rows.append(dict(row))

            return headers, data_rows

    def _preview_excel_data(
        self, file_path: str, max_rows: int
    ) -> tuple[list[str], list[dict[str, Any]]]:
        """预览Excel数据"""
        import pandas as pd

        df = pd.read_excel(file_path, nrows=max_rows)
        headers = df.columns.tolist()
        data_rows = df.to_dict("records")

        return headers, data_rows

    def import_data(
        self,
        file_path: str,
        data_type: str,
        field_mapping: dict[str, str],
        options: dict[str, Any] | None = None,
    ) -> tuple[int, int, list[str]]:
        """
        导入数据

        Args:
            file_path: 文件路径
            data_type: 数据类型 (customers, suppliers, contracts)
            field_mapping: 字段映射 {目标字段: 源字段}
            options: 导入选项

        Returns:
            Tuple[int, int, List[str]]: (成功数量, 失败数量, 错误信息列表)

        Raises:
            ServiceError: 导入失败时抛出
        """
        try:
            self._logger.info(f"开始导入数据: {data_type} from {file_path}")

            # 验证数据类型
            if data_type not in self._data_type_services:
                raise ServiceError(f"不支持的数据类型: {data_type}")

            # 读取文件数据
            file_ext = Path(file_path).suffix.lower()
            if file_ext == ".csv":
                raw_data = self._read_csv_data(file_path)
            elif file_ext == ".xlsx":
                raw_data = self._read_excel_data(file_path)
            else:
                raise ServiceError(f"不支持的文件格式: {file_ext}")

            # 映射字段
            mapped_data = self._map_fields(raw_data, field_mapping)

            # 导入数据
            return self._import_mapped_data(data_type, mapped_data, options or {})

        except Exception as e:
            self._logger.error(f"导入数据失败: {e}")
            raise ServiceError(f"导入数据失败: {e}") from e

    def _read_csv_data(self, file_path: str) -> list[dict[str, Any]]:
        """读取CSV数据"""
        with open(file_path, encoding="utf-8-sig") as file:
            sample = file.read(1024)
            file.seek(0)
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter

            reader = csv.DictReader(file, delimiter=delimiter)
            return list(reader)

    def _read_excel_data(self, file_path: str) -> list[dict[str, Any]]:
        """读取Excel数据"""
        import pandas as pd

        df = pd.read_excel(file_path)
        # 处理NaN值
        df = df.fillna("")
        return df.to_dict("records")

    def _map_fields(
        self, raw_data: list[dict[str, Any]], field_mapping: dict[str, str]
    ) -> list[dict[str, Any]]:
        """映射字段"""
        mapped_data = []

        for row in raw_data:
            mapped_row = {}
            for target_field, source_field in field_mapping.items():
                if source_field and source_field in row:
                    mapped_row[target_field] = row[source_field]

            # 只保留有效数据行
            if any(mapped_row.values()):
                mapped_data.append(mapped_row)

        return mapped_data

    def _import_mapped_data(
        self, data_type: str, mapped_data: list[dict[str, Any]], options: dict[str, Any]
    ) -> tuple[int, int, list[str]]:
        """导入映射后的数据"""
        service = self._data_type_services[data_type]
        success_count = 0
        error_count = 0
        error_messages = []

        # TODO: 实现重复数据处理和更新逻辑
        # skip_duplicates = options.get("skip_duplicates", True)
        # update_existing = options.get("update_existing", False)

        for i, row_data in enumerate(mapped_data):
            try:
                # 数据验证
                if data_type == "customers":
                    validation_result = validate_customer_data(row_data)
                elif data_type == "suppliers":
                    validation_result = validate_supplier_data(row_data)
                else:
                    # 对于其他类型,进行基本验证
                    validation_result = self._basic_validation(row_data)

                if not validation_result.is_valid:
                    errors_str = ", ".join(validation_result.errors)
                    error_msg = f"第{i + 1}行数据验证失败: {errors_str}"
                    error_messages.append(error_msg)
                    error_count += 1
                    continue

                # 创建记录
                record_id: int | None = None
                if data_type == "customers":
                    record_id = self._customer_service.create_customer(row_data)
                elif data_type == "suppliers":
                    # 检查供应商服务是否有create_supplier方法
                    if hasattr(self._supplier_service, "create_supplier"):
                        record_id = self._supplier_service.create_supplier(row_data)
                    else:
                        # 使用通用创建方法
                        record_id = self._create_other_record(service, row_data)
                else:
                    # 其他类型的创建逻辑
                    record_id = self._create_other_record(service, row_data)

                if record_id:
                    success_count += 1
                else:
                    error_count += 1
                    error_messages.append(f"第{i + 1}行数据创建失败")

            except Exception as e:
                error_count += 1
                error_messages.append(f"第{i + 1}行处理失败: {e}")

        self._logger.info(f"数据导入完成: 成功{success_count}条, 失败{error_count}条")
        return success_count, error_count, error_messages

    def _basic_validation(self, data: dict[str, Any]) -> Any:
        """基本数据验证"""
        from transfunctions.validation import ValidationResult

        errors = []

        # 检查必要字段
        if not data.get("name"):
            errors.append("名称不能为空")

        return ValidationResult(is_valid=len(errors) == 0, errors=errors, warnings=[])

    def _create_other_record(self, service: Any, data: dict[str, Any]) -> int | None:
        """创建其他类型记录的通用方法"""
        # 这里可以根据具体的服务类型实现不同的创建逻辑
        if hasattr(service, "create"):
            return service.create(data)
        return None

    def export_data(
        self,
        data_type: str,
        export_format: str,
        output_path: str,
        filters: dict[str, Any] | None = None,
        fields: list[str] | None = None,
    ) -> bool:
        """
        导出数据

        Args:
            data_type: 数据类型
            export_format: 导出格式 (.csv, .xlsx, .pdf)
            output_path: 输出路径
            filters: 筛选条件
            fields: 导出字段列表

        Returns:
            bool: 导出是否成功

        Raises:
            ServiceError: 导出失败时抛出
        """
        try:
            self._logger.info(f"开始导出数据: {data_type} to {output_path}")

            # 获取数据
            data = self._get_export_data(data_type, filters)

            if not data:
                raise ServiceError("没有数据可以导出")

            # 筛选字段
            if fields:
                data = self._filter_fields(data, fields)

            # 根据格式导出
            if export_format == ".csv":
                return self._export_csv(data, output_path)
            elif export_format == ".xlsx":
                return self._export_excel(data, output_path)
            elif export_format == ".pdf":
                return self._export_pdf(data, output_path, data_type)
            else:
                raise ServiceError(f"不支持的导出格式: {export_format}")

        except Exception as e:
            self._logger.error(f"导出数据失败: {e}")
            raise ServiceError(f"导出数据失败: {e}") from e

    def _get_export_data(
        self, data_type: str, filters: dict[str, Any] | None
    ) -> list[dict[str, Any]]:
        """获取导出数据"""
        service = self._data_type_services.get(data_type)
        if not service:
            raise ServiceError(f"不支持的数据类型: {data_type}")

        # 根据数据类型获取数据
        if data_type == "customers":
            customers, _ = self._customer_service.search_customers(
                query=filters.get("query", "") if filters else "",
                filters=filters.get("filters") if filters else None,
                page=1,
                page_size=10000,  # 大批量导出
            )
            return customers
        elif data_type == "suppliers":
            # 假设供应商服务有类似的搜索方法
            if hasattr(service, "search_suppliers"):
                suppliers, _ = service.search_suppliers(
                    query=filters.get("query", "") if filters else "",
                    filters=filters.get("filters") if filters else None,
                    page=1,
                    page_size=10000,
                )
                return suppliers
            else:
                return (
                    service.get_all_suppliers()
                    if hasattr(service, "get_all_suppliers")
                    else []
                )
        else:
            # 其他数据类型的获取逻辑
            return []

    def _filter_fields(
        self, data: list[dict[str, Any]], fields: list[str]
    ) -> list[dict[str, Any]]:
        """筛选字段"""
        filtered_data = []
        for row in data:
            filtered_row = {field: row.get(field, "") for field in fields}
            filtered_data.append(filtered_row)
        return filtered_data

    def _export_csv(self, data: list[dict[str, Any]], output_path: str) -> bool:
        """导出CSV格式"""
        if not data:
            return False

        with open(output_path, "w", newline="", encoding="utf-8-sig") as csvfile:
            fieldnames = list(data[0].keys())
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)

        return True

    def _export_excel(self, data: list[dict[str, Any]], output_path: str) -> bool:
        """导出Excel格式"""
        try:
            import pandas as pd

            df = pd.DataFrame(data)
            df.to_excel(output_path, index=False, engine="openpyxl")
            return True

        except ImportError as e:
            raise ServiceError("导出Excel需要安装pandas和openpyxl库") from e

    def _export_pdf(
        self, data: list[dict[str, Any]], output_path: str, data_type: str
    ) -> bool:
        """导出PDF格式"""
        # 这里可以根据数据类型生成不同的PDF报表
        # 暂时使用简单的表格格式
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

            doc = SimpleDocTemplate(output_path, pagesize=A4)
            elements = []

            if data:
                # 创建表格数据
                headers = list(data[0].keys())
                table_data = [headers]

                table_data.extend(
                    [str(row.get(field, "")) for field in headers] for row in data
                )

                # 创建表格
                table = Table(table_data)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 14),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )

                elements.append(table)

            doc.build(elements)
            return True

        except Exception as e:
            self._logger.error(f"PDF导出失败: {e}")
            return False

    def generate_word_document(
        self, template_type: str, data: dict[str, Any], output_path: str
    ) -> bool:
        """
        生成Word文档

        Args:
            template_type: 模板类型 (contract, quote, customer_report)
            data: 数据字典
            output_path: 输出路径

        Returns:
            bool: 生成是否成功
        """
        try:
            self._logger.info(f"生成Word文档: {template_type} -> {output_path}")

            # 简单的Word文档生成实现
            from docx import Document

            doc = Document()
            doc.add_heading(f"{template_type.title()} 文档", 0)

            # 添加数据内容
            for key, value in data.items():
                doc.add_paragraph(f"{key}: {value}")

            doc.save(output_path)
            return True

        except ImportError:
            self._logger.error("生成Word文档需要安装python-docx库")
            raise ServiceError("生成Word文档需要安装python-docx库")
        except Exception as e:
            self._logger.error(f"生成Word文档失败: {e}")
            raise ServiceError(f"生成Word文档失败: {e}") from e

    def generate_pdf_document(
        self, template_type: str, data: dict[str, Any], output_path: str
    ) -> bool:
        """
        生成PDF文档

        Args:
            template_type: 模板类型 (contract, quote, customer_report)
            data: 数据字典
            output_path: 输出路径

        Returns:
            bool: 生成是否成功
        """
        try:
            self._logger.info(f"生成PDF文档: {template_type} -> {output_path}")

            # 使用reportlab直接生成PDF
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # 添加标题
            title = Paragraph(f"{template_type.title()} 文档", styles["Title"])
            story.append(title)
            story.append(Spacer(1, 12))

            # 添加数据内容
            for key, value in data.items():
                content = Paragraph(f"<b>{key}:</b> {value}", styles["Normal"])
                story.append(content)
                story.append(Spacer(1, 6))

            doc.build(story)
            return True

        except ImportError:
            self._logger.error("生成PDF文档需要安装reportlab库")
            raise ServiceError("生成PDF文档需要安装reportlab库")
        except Exception as e:
            self._logger.error(f"生成PDF文档失败: {e}")
            raise ServiceError(f"生成PDF文档失败: {e}") from e

    def get_available_templates(self) -> list[str]:
        """
        获取可用的文档模板

        Returns:
            List[str]: 可用模板类型列表
        """
        return ["contract", "quote", "customer_report"]

    def get_field_mapping_suggestions(
        self, data_type: str, source_fields: list[str]
    ) -> dict[str, str]:
        """
        获取字段映射建议

        Args:
            data_type: 数据类型
            source_fields: 源字段列表

        Returns:
            Dict[str, str]: 映射建议 {目标字段: 源字段}
        """
        # 定义字段映射规则
        mapping_rules = {
            "customers": {
                "name": ["name", "客户名称", "公司名称", "名称", "company_name"],
                "contact_person": ["contact", "联系人", "姓名", "contact_person"],
                "phone": ["phone", "电话", "手机", "联系电话", "telephone"],
                "email": ["email", "邮箱", "邮件", "电子邮件"],
                "company": ["company", "公司", "企业", "单位"],
                "address": ["address", "地址", "联系地址"],
                "notes": ["notes", "备注", "说明", "描述", "remark"],
            },
            "suppliers": {
                "name": ["name", "供应商名称", "公司名称", "名称"],
                "contact_person": ["contact", "联系人", "姓名"],
                "phone": ["phone", "电话", "手机", "联系电话"],
                "email": ["email", "邮箱", "邮件"],
                "company": ["company", "公司", "企业"],
                "address": ["address", "地址"],
                "notes": ["notes", "备注", "说明"],
            },
        }

        suggestions = {}
        rules = mapping_rules.get(data_type, {})

        for target_field, patterns in rules.items():
            for pattern in patterns:
                for source_field in source_fields:
                    if pattern.lower() in source_field.lower():
                        suggestions[target_field] = source_field
                        break
                if target_field in suggestions:
                    break

        return suggestions
