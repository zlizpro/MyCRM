"""
MiniCRM Word文档生成服务

专门负责Word文档的生成,包括:
- 基于模板的Word文档生成
- 简单Word文档创建
- 文档内容填充和格式化
- Word文档验证

设计原则:
- 单一职责:只负责Word文档生成
- 支持模板和非模板生成
- 完整的错误处理和日志记录
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from minicrm.core.exceptions import ServiceError


class WordDocumentService:
    """
    Word文档生成服务

    负责生成各种Word文档,支持模板和非模板方式.
    """

    def __init__(self):
        """初始化Word文档生成服务"""
        self._logger = logging.getLogger(__name__)
        self._logger.info("Word文档生成服务初始化完成")

    def generate_from_template(
        self, template_path: Path, data: dict[str, Any], output_path: str
    ) -> bool:
        """
        基于模板生成Word文档

        Args:
            template_path: 模板文件路径
            data: 数据字典
            output_path: 输出文件路径

        Returns:
            bool: 生成是否成功

        Raises:
            ServiceError: 生成失败时抛出
        """
        try:
            self._logger.info(f"基于模板生成Word文档: {template_path} -> {output_path}")

            # 检查模板文件是否存在
            if not template_path.exists():
                raise ServiceError(f"模板文件不存在: {template_path}")

            # 尝试使用docxtpl处理模板
            try:
                from docxtpl import DocxTemplate

                # 加载模板
                doc = DocxTemplate(str(template_path))

                # 准备数据
                context = self._prepare_template_data(data)

                # 渲染模板
                doc.render(context)

                # 保存文档
                doc.save(output_path)

                self._logger.info(f"基于模板的Word文档生成成功: {output_path}")
                return True

            except ImportError:
                # 如果没有docxtpl库,使用简单方式生成
                self._logger.warning("docxtpl库未安装,使用简单方式生成文档")
                return self.generate_simple_document(data, output_path)

            except Exception as e:
                # 如果模板处理失败,使用简单方式生成
                self._logger.warning(f"模板处理失败,使用简单方式生成: {e}")
                return self.generate_simple_document(data, output_path)

        except Exception as e:
            self._logger.error(f"基于模板生成Word文档失败: {e}")
            raise ServiceError(f"生成Word文档失败: {e}") from e

    def generate_simple_document(
        self, data: dict[str, Any], output_path: str, document_type: str = "generic"
    ) -> bool:
        """
        生成简单的Word文档(不使用模板)

        Args:
            data: 数据字典
            output_path: 输出路径
            document_type: 文档类型

        Returns:
            bool: 创建是否成功
        """
        try:
            from docx import Document

            # 创建新文档
            doc = Document()

            # 根据文档类型生成不同内容
            if document_type == "contract":
                self._create_contract_content(doc, data)
            elif document_type == "quote":
                self._create_quote_content(doc, data)
            elif document_type == "customer_report":
                self._create_customer_report_content(doc, data)
            else:
                self._create_generic_content(doc, data)

            # 保存文档
            doc.save(output_path)

            self._logger.info(f"简单Word文档生成成功: {output_path}")
            return True

        except ImportError:
            # 如果连python-docx都没有,创建文本文件
            self._logger.warning("python-docx库未安装,生成文本文件")
            return self._create_text_document(data, output_path, document_type)
        except Exception as e:
            self._logger.error(f"创建简单Word文档失败: {e}")
            return False

    def _create_contract_content(self, doc, data: dict[str, Any]) -> None:
        """创建合同内容"""
        # 标题
        title = doc.add_heading("合同文档", 0)
        title.alignment = 1  # 居中

        # 合同信息
        doc.add_heading("合同信息", level=1)

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = "Table Grid"

        # 填充合同信息
        contract_info = [
            ("合同编号", data.get("contract_number", "")),
            ("客户名称", data.get("customer_name", "")),
            ("合同金额", data.get("contract_amount", "")),
            ("签署日期", data.get("sign_date", "")),
            ("生效日期", data.get("effective_date", "")),
            ("到期日期", data.get("expire_date", "")),
        ]

        for i, (label, value) in enumerate(contract_info):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)

        # 合同条款
        doc.add_heading("合同条款", level=1)
        content = data.get("contract_content", "合同条款内容...")
        doc.add_paragraph(content)

        # 签署信息
        doc.add_heading("签署信息", level=1)
        doc.add_paragraph("甲方(客户):_________________")
        doc.add_paragraph("乙方(公司):_________________")
        doc.add_paragraph(f"签署日期:{datetime.now().strftime('%Y年%m月%d日')}")

    def _create_quote_content(self, doc, data: dict[str, Any]) -> None:
        """创建报价单内容"""
        # 标题
        title = doc.add_heading("产品报价单", 0)
        title.alignment = 1  # 居中

        # 报价信息
        doc.add_heading("报价信息", level=1)

        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"

        quote_info = [
            ("报价编号", data.get("quote_number", "")),
            ("客户名称", data.get("customer_name", "")),
            ("报价日期", data.get("quote_date", "")),
            ("有效期至", data.get("valid_until", "")),
        ]

        for i, (label, value) in enumerate(quote_info):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)

        # 产品清单
        doc.add_heading("产品清单", level=1)

        items = data.get("items", [])
        if items:
            # 创建产品表格
            items_table = doc.add_table(rows=len(items) + 1, cols=5)
            items_table.style = "Table Grid"

            # 表头
            headers = ["产品名称", "规格型号", "数量", "单价", "小计"]
            for i, header in enumerate(headers):
                items_table.cell(0, i).text = header

            # 产品数据
            total_amount = 0
            for i, item in enumerate(items):
                items_table.cell(i + 1, 0).text = item.get("name", "")
                items_table.cell(i + 1, 1).text = item.get("model", "")
                items_table.cell(i + 1, 2).text = str(item.get("quantity", ""))
                items_table.cell(i + 1, 3).text = str(item.get("price", ""))
                subtotal = item.get("subtotal", 0)
                items_table.cell(i + 1, 4).text = str(subtotal)
                total_amount += float(subtotal) if subtotal else 0
        else:
            doc.add_paragraph("暂无产品项目")
            total_amount = data.get("total_amount", 0)

        # 总计
        doc.add_paragraph(f"总计金额:¥{total_amount:,.2f}")

        # 备注
        doc.add_heading("备注", level=1)
        notes = data.get("notes", "本报价单有效期为30天,如有疑问请及时联系.")
        doc.add_paragraph(notes)

    def _create_customer_report_content(self, doc, data: dict[str, Any]) -> None:
        """创建客户报告内容"""
        # 标题
        title_text = data.get("report_title", "客户分析报告")
        title = doc.add_heading(title_text, 0)
        title.alignment = 1  # 居中

        # 报告信息
        doc.add_paragraph(
            f"报告日期:{data.get('report_date', datetime.now().strftime('%Y-%m-%d'))}"
        )
        doc.add_paragraph(f"客户名称:{data.get('customer_name', '')}")

        # 客户基本信息
        doc.add_heading("客户基本信息", level=1)
        customer_info = data.get("customer_info", {})
        if customer_info:
            for key, value in customer_info.items():
                doc.add_paragraph(f"{key}:{value}")
        else:
            doc.add_paragraph("客户基本信息...")

        # 互动历史
        doc.add_heading("互动历史", level=1)
        interaction_history = data.get("interaction_history", [])
        if interaction_history:
            for interaction in interaction_history:
                doc.add_paragraph(f"• {interaction}")
        else:
            doc.add_paragraph("暂无互动记录")

        # 财务摘要
        doc.add_heading("财务摘要", level=1)
        financial_summary = data.get("financial_summary", {})
        if financial_summary:
            for key, value in financial_summary.items():
                doc.add_paragraph(f"{key}:{value}")
        else:
            doc.add_paragraph("财务摘要信息...")

    def _create_generic_content(self, doc, data: dict[str, Any]) -> None:
        """创建通用内容"""
        # 标题
        title_text = data.get("title", "文档")
        title = doc.add_heading(title_text, 0)
        title.alignment = 1  # 居中

        # 内容
        content = data.get("content", "文档内容...")
        doc.add_paragraph(content)

        # 生成时间
        doc.add_paragraph(
            f"生成时间:{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        )

    def _create_text_document(
        self, data: dict[str, Any], output_path: str, document_type: str
    ) -> bool:
        """
        创建文本文档(备用方案)

        Args:
            data: 数据字典
            output_path: 输出路径
            document_type: 文档类型

        Returns:
            bool: 创建是否成功
        """
        try:
            # 将.docx扩展名改为.txt
            if output_path.endswith(".docx"):
                output_path = output_path.replace(".docx", ".txt")

            content = self._generate_text_content(document_type, data)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

            self._logger.info(f"文本文档生成成功: {output_path}")
            return True

        except Exception as e:
            self._logger.error(f"创建文本文档失败: {e}")
            return False

    def _generate_text_content(self, document_type: str, data: dict[str, Any]) -> str:
        """生成文本内容"""
        if document_type == "contract":
            return self._generate_contract_text(data)
        elif document_type == "quote":
            return self._generate_quote_text(data)
        elif document_type == "customer_report":
            return self._generate_customer_report_text(data)
        else:
            return self._generate_generic_text(data)

    def _generate_contract_text(self, data: dict[str, Any]) -> str:
        """生成合同文本内容"""
        return f"""
合同文档
========

合同编号:{data.get("contract_number", "")}
客户名称:{data.get("customer_name", "")}
合同金额:{data.get("contract_amount", "")}
签署日期:{data.get("sign_date", "")}
生效日期:{data.get("effective_date", "")}
到期日期:{data.get("expire_date", "")}

合同条款:
{data.get("contract_content", "合同条款内容...")}

签署信息:
甲方(客户):_________________
乙方(公司):_________________
签署日期:{datetime.now().strftime("%Y年%m月%d日")}
"""

    def _generate_quote_text(self, data: dict[str, Any]) -> str:
        """生成报价单文本内容"""
        items_text = self._format_quote_items_text(data.get("items", []))

        return f"""
产品报价单
==========

报价编号:{data.get("quote_number", "")}
客户名称:{data.get("customer_name", "")}
报价日期:{data.get("quote_date", "")}
有效期至:{data.get("valid_until", "")}

产品清单:
{items_text}

总计金额:¥{data.get("total_amount", 0)}

备注:
{data.get("notes", "本报价单有效期为30天,如有疑问请及时联系.")}
"""

    def _generate_customer_report_text(self, data: dict[str, Any]) -> str:
        """生成客户报告文本内容"""
        customer_info_text = self._format_customer_info_text(
            data.get("customer_info", {})
        )
        interaction_text = self._format_interaction_history_text(
            data.get("interaction_history", [])
        )
        financial_text = self._format_financial_summary_text(
            data.get("financial_summary", {})
        )

        return f"""
{data.get("report_title", "客户分析报告")}
{"=" * 20}

报告日期:{data.get("report_date", datetime.now().strftime("%Y-%m-%d"))}
客户名称:{data.get("customer_name", "")}

客户基本信息:
{customer_info_text}

互动历史:
{interaction_text}

财务摘要:
{financial_text}
"""

    def _generate_generic_text(self, data: dict[str, Any]) -> str:
        """生成通用文本内容"""
        return f"""
{data.get("title", "文档")}
{"=" * 20}

{data.get("content", "文档内容...")}

生成时间:{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}
"""

    def _format_quote_items_text(self, items: list[dict[str, Any]]) -> str:
        """格式化报价项目文本"""
        if not items:
            return "暂无产品项目"

        text = "产品名称\t规格型号\t数量\t单价\t小计\n"
        text += "-" * 50 + "\n"

        for item in items:
            name = item.get("name", "")
            model = item.get("model", "")
            quantity = item.get("quantity", "")
            price = item.get("price", "")
            subtotal = item.get("subtotal", "")
            text += f"{name}\t{model}\t{quantity}\t{price}\t{subtotal}\n"

        return text

    def _format_customer_info_text(self, info: dict[str, Any]) -> str:
        """格式化客户信息文本"""
        if not info:
            return "客户基本信息..."

        text = ""
        for key, value in info.items():
            text += f"{key}:{value}\n"

        return text

    def _format_interaction_history_text(self, history: list[str]) -> str:
        """格式化互动历史文本"""
        if not history:
            return "暂无互动记录"

        text = ""
        for interaction in history:
            text += f"• {interaction}\n"

        return text

    def _format_financial_summary_text(self, summary: dict[str, Any]) -> str:
        """格式化财务摘要文本"""
        if not summary:
            return "财务摘要信息..."

        text = ""
        for key, value in summary.items():
            text += f"{key}:{value}\n"

        return text

    def _prepare_template_data(self, data: dict[str, Any]) -> dict[str, Any]:
        """
        准备模板数据

        Args:
            data: 原始数据

        Returns:
            Dict[str, Any]: 处理后的模板数据
        """
        # 添加通用数据
        context = {
            "current_date": datetime.now().strftime("%Y年%m月%d日"),
            "current_time": datetime.now().strftime("%H:%M:%S"),
            **data,
        }

        # 处理产品项目(如果存在)
        items = context.get("items", [])
        if items:
            total = sum(float(item.get("subtotal", 0)) for item in items)
            context["total_amount"] = f"¥{total:,.2f}"

        return context
